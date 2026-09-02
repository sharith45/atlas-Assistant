import os
import sys
import json
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Configurar encoding en Windows para prevenir UnicodeEncodeError en la consola
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

def hex_to_rgb(hex_str):
    """Convierte un color en hexadecimal (con o sin #) a RGBColor de python-docx."""
    hex_str = hex_str.lstrip('#')
    if len(hex_str) == 3:
        hex_str = ''.join([c*2 for c in hex_str])
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece padding interno en celdas de Word (en dxa: 1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_border(paragraph, color_hex="1B365D", size=12, border_type="bottom"):
    """Agrega un borde sutil de color superior o inferior a un párrafo para división visual."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    
    border = OxmlElement(f'w:{border_type}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size)) # 1/8 pt, ej. 12 = 1.5 pt
    border.set(qn('w:space'), '6')
    border.set(qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(border)

def add_watermark(doc, text="CONFIDENCIAL"):
    """Inserta una marca de agua de texto en color gris claro en el encabezado de las secciones."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Eliminar cualquier run previo en el header
        p.text = ""
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(36)
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

def add_cover_page(doc, cover_data):
    """Genera una portada formal con jerarquía de título, subtítulo, autor y línea decorativa."""
    # Saltos de línea para bajar el título
    for _ in range(4):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = cover_data.get("title", "Título del Documento").upper()
    title_run = title_p.add_run(title_text)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.bold = True
    color_hex = cover_data.get("theme_color", "1B365D")
    title_run.font.color.rgb = hex_to_rgb(color_hex)
    
    # Línea horizontal divisoria debajo del título
    add_paragraph_border(title_p, color_hex=color_hex, size=18, border_type="bottom")
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(14)
    sub_run = subtitle_p.add_run(cover_data.get("subtitle", ""))
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    for _ in range(7):
        doc.add_paragraph()
        
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author = cover_data.get("author", "")
    date = cover_data.get("date", "")
    org = cover_data.get("organization", "")
    
    info_text = ""
    if author: info_text += f"Elaborado por: {author}\n"
    if org: info_text += f"Organización: {org}\n"
    if date: info_text += f"Fecha: {date}"
    
    info_run = info_p.add_run(info_text.strip())
    info_run.font.name = 'Arial'
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    doc.add_page_break()

def add_signature_block(doc, signatures, color_hex="1B365D"):
    """Inserta una sección estructurada de firmas alineada al final del documento."""
    num_sigs = len(signatures)
    if num_sigs == 0:
        return
        
    table = doc.add_table(rows=2, cols=num_sigs)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Quitar bordes para que sea invisible
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # Fila 0: Línea para firmar
    for idx, sig in enumerate(signatures):
        cell = table.rows[0].cells[idx]
        cell.text = "_________________________"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(2)
        if len(p.runs) > 0:
            p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            
    # Fila 1: Nombres y Roles
    for idx, sig in enumerate(signatures):
        cell = table.rows[1].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        
        run_name = p.add_run(f"{sig.get('name', '')}\n")
        run_name.font.name = 'Arial'
        run_name.bold = True
        run_name.font.size = Pt(10)
        run_name.font.color.rgb = hex_to_rgb(color_hex)
        
        run_role = p.add_run(sig.get('role', ''))
        run_role.font.name = 'Arial'
        run_role.font.size = Pt(9)
        run_role.italic = True
        run_role.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

def build_docx_from_json(json_data):
    """Construye un documento de Word completo a partir del JSON estructurado."""
    doc = Document()
    
    # 1. Ajustes globales
    settings = json_data.get("settings", {})
    margin_val = settings.get("margin_inch", 1.0)
    
    for section in doc.sections:
        section.top_margin = Inches(margin_val)
        section.bottom_margin = Inches(margin_val)
        section.left_margin = Inches(margin_val)
        section.right_margin = Inches(margin_val)
        
        if settings.get("orientation") == "landscape":
            # Cambiar dimensiones para apaisado
            section.orientation = section.orientation # Reset
            h, w = section.page_height, section.page_width
            section.page_width = max(h, w)
            section.page_height = min(h, w)
            
    # Configuración de tipografía base
    font_name = settings.get("font_name", "Arial")
    base_color = settings.get("theme_color", "1B365D")
    
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = font_name
    normal_font.size = Pt(settings.get("font_size", 11))
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = settings.get("line_spacing", 1.15)
    normal_style.paragraph_format.space_after = Pt(6)
    
    # 2. Marca de agua
    watermark_text = settings.get("watermark")
    if watermark_text:
        add_watermark(doc, watermark_text)
        
    # 3. Procesar elementos secuencialmente
    elements = json_data.get("elements", [])
    for element in elements:
        elem_type = element.get("type")
        
        if elem_type == "cover_page":
            add_cover_page(doc, {**element, "theme_color": base_color})
            
        elif elem_type == "heading":
            level = element.get("level", 1)
            text = element.get("text", "")
            heading = doc.add_heading(text, level=level)
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(6)
            
            # Línea decorativa superior/inferior si se solicita
            border_pos = element.get("border") # "top" o "bottom"
            if border_pos:
                add_paragraph_border(heading, color_hex=base_color, size=12, border_type=border_pos)
            
            if len(heading.runs) > 0:
                run = heading.runs[0]
                run.font.name = font_name
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                    run.font.color.rgb = hex_to_rgb(base_color)
                elif level == 2:
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0x4A, 0x6B, 0x82)
                else:
                    run.font.size = Pt(11.5)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    
        elif elem_type == "paragraph":
            text = element.get("text", "")
            style = element.get("style", "Normal")
            align = element.get("align", "left")
            
            p = doc.add_paragraph()
            if style == "bullet":
                p.style = 'List Bullet'
                
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            run = p.add_run(text)
            run.font.name = font_name
            run.bold = element.get("bold", False)
            run.italic = element.get("italic", False)
            
            # Color del párrafo si está configurado
            p_color = element.get("color")
            if p_color:
                run.font.color.rgb = hex_to_rgb(p_color)
                
            # Línea divisoria en párrafo
            border_pos = element.get("border")
            if border_pos:
                add_paragraph_border(p, color_hex=base_color, size=6, border_type=border_pos)
                
        elif elem_type == "table":
            headers = element.get("headers", [])
            rows = element.get("rows", [])
            widths = element.get("widths_inches", [])
            align = element.get("align", "center")
            
            num_rows = len(rows) + (1 if headers else 0)
            num_cols = max(len(headers), max((len(r) for r in rows), default=0)) if rows or headers else 0
            
            if num_rows > 0 and num_cols > 0:
                table = doc.add_table(rows=num_rows, cols=num_cols)
                if align == "center":
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif align == "left":
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                # Aplicar bordes
                tblPr = table._tbl.tblPr
                borders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>\n'
                    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{base_color}"/>\n'
                    f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
                    f'  <w:insideV w:val="none"/>\n'
                    f'  <w:left w:val="none"/>\n'
                    f'  <w:right w:val="none"/>\n'
                    f'</w:tblBorders>'
                )
                tblPr.append(borders)
                
                # Asignar anchos de columnas
                for row in table.rows:
                    for i, w_val in enumerate(widths):
                        if i < len(row.cells):
                            row.cells[i].width = Inches(w_val)
                            
                current_row_idx = 0
                if headers:
                    hdr_cells = table.rows[0].cells
                    for i, h_text in enumerate(headers):
                        if i < len(hdr_cells):
                            hdr_cells[i].text = str(h_text)
                            # Fondo color del tema
                            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{base_color}"/>')
                            tcPr.append(shd)
                            set_cell_margins(hdr_cells[i])
                            
                            p = hdr_cells[i].paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.runs[0]
                            run.font.name = font_name
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            
                    trPr = table.rows[0]._tr.get_or_add_trPr()
                    trPr.append(OxmlElement('w:cantSplit'))
                    trPr.append(OxmlElement('w:tblHeader'))
                    current_row_idx = 1
                    
                # Rellenar datos
                for r_data in rows:
                    row_cells = table.rows[current_row_idx].cells
                    trPr = table.rows[current_row_idx]._tr.get_or_add_trPr()
                    trPr.append(OxmlElement('w:cantSplit'))
                    
                    for i, cell_content in enumerate(r_data):
                        if i < len(row_cells):
                            cell = row_cells[i]
                            
                            # Comprobar si la celda viene formateada como dict
                            if isinstance(cell_content, dict):
                                cell.text = str(cell_content.get("value", ""))
                                set_cell_margins(cell)
                                
                                fill = cell_content.get("fill")
                                if fill:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
                                    tcPr.append(shd)
                                    
                                p = cell.paragraphs[0]
                                c_align = cell_content.get("align", "left")
                                if c_align == "right":
                                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                elif c_align == "center":
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                if len(p.runs) > 0:
                                    run = p.runs[0]
                                    run.font.name = font_name
                                    run.font.size = Pt(9.5)
                                    run.bold = cell_content.get("bold", False)
                                    run.italic = cell_content.get("italic", False)
                                    c_color = cell_content.get("color")
                                    if c_color:
                                        run.font.color.rgb = hex_to_rgb(c_color)
                            else:
                                # Celda de texto simple
                                cell.text = str(cell_content)
                                set_cell_margins(cell)
                                p = cell.paragraphs[0]
                                if len(p.runs) > 0:
                                    run = p.runs[0]
                                    run.font.name = font_name
                                    run.font.size = Pt(9.5)
                                    
                    current_row_idx += 1
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                
        elif elem_type == "image":
            path = element.get("path")
            width_val = element.get("width_inches", 4.0)
            align = element.get("align", "center")
            caption = element.get("caption")
            
            if path and os.path.exists(path):
                p = doc.add_paragraph()
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                p.add_run().add_picture(path, width=Inches(width_val))
                
                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(4)
                    run_cap = p_cap.add_run(caption)
                    run_cap.font.name = font_name
                    run_cap.font.size = Pt(9)
                    run_cap.italic = True
                    run_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            else:
                print(f"⚠️ Warning: Imagen no encontrada en la ruta '{path}'")
                
        elif elem_type == "page_break":
            doc.add_page_break()
            
        elif elem_type == "signature_block":
            add_signature_block(doc, element.get("signatures", []), base_color)
            
    return doc

def main():
    parser = argparse.ArgumentParser(
        description="Generador automatizado de documentos Word (.docx) profesionales a partir de especificaciones JSON."
    )
    parser.add_argument("input_json", help="Ruta al archivo JSON de entrada con la especificación")
    parser.add_argument("output_docx", nargs="?", help="Ruta de salida para el documento de Word (opcional)")
    args = parser.parse_args()
    
    input_json = os.path.abspath(args.input_json)
    if not os.path.exists(input_json):
        print(f"❌ Error: Archivo JSON no encontrado: {input_json}", file=sys.stderr)
        sys.exit(1)
        
    if args.output_docx:
        output_docx = os.path.abspath(args.output_docx)
    else:
        output_docx = os.path.splitext(input_json)[0] + ".docx"
        
    # Crear carpeta de salida si no existe
    out_dir = os.path.dirname(output_docx)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
        
    print(f"📄 Leyendo especificación desde: {input_json}")
    
    try:
        with open(input_json, "r", encoding="utf-8") as f:
            json_data = json.load(f)
            
        doc = build_docx_from_json(json_data)
        doc.save(output_docx)
        
        size = os.path.getsize(output_docx)
        print(f"✅ [OK] Documento Word generado exitosamente: {output_docx} ({size / 1024:.1f} KB)")
        sys.exit(0)
    except Exception as e:
        import traceback
        print(f"❌ [ERROR] Falló la compilación del documento: {e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
